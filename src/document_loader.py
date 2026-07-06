"""
文档解析模块
支持 Word(.docx)、PDF、TXT 三种格式
"""
import os
import re
import chardet
from typing import Dict, Optional, Tuple


def _is_text_garbled(text: str) -> Tuple[bool, float]:
    """
    判断文本是否可能是乱码/二进制内容。
    返回 (is_garbled, printable_ratio)
    """
    if not text:
        return True, 0.0
    # 中文字符、常见标点、字母数字、空白
    # 中文字符、常见标点、字母数字、空白
    printable = len(re.findall(
        r'[一-鿿　-〿＀-￯'
        r'a-zA-Z0-9\s.,;:!?()、。，；：！？'
        r'“”《》【】—…+\-*/=|#@%&]',
        text))
    ratio = printable / len(text)
    return ratio < 0.5, ratio


def _detect_encoding(filepath: str) -> str:
    """自动检测文件编码"""
    with open(filepath, "rb") as f:
        raw = f.read(10000)  # 读前 10KB 即可检测
    result = chardet.detect(raw)
    return result.get("encoding", "utf-8") or "utf-8"


def load_txt(filepath: str) -> Dict:
    """
    加载 TXT 文件，自动检测编码（UTF-8/GBK/GB2312）
    返回 {"filename": ..., "text": ..., "source_type": "txt"}
    """
    encoding = _detect_encoding(filepath)
    with open(filepath, "r", encoding=encoding, errors="replace") as f:
        text = f.read()
    return {
        "filename": os.path.basename(filepath),
        "text": text,
        "source_type": "txt",
    }


def load_word(filepath: str) -> Dict:
    """
    加载 .docx 文件，提取段落和表格文本
    返回 {"filename": ..., "text": ..., "source_type": "docx"}
    """
    from docx import Document

    doc = Document(filepath)
    parts = []

    # 提取段落文本
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # 提取表格文本
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                parts.append(" | ".join(row_text))

    return {
        "filename": os.path.basename(filepath),
        "text": "\n".join(parts),
        "source_type": "docx",
    }


def load_pdf(filepath: str) -> Dict:
    """
    加载 PDF 文件，提取全部文本
    返回 {"filename": ..., "text": ..., "source_type": "pdf"}
    """
    import fitz  # PyMuPDF

    doc = fitz.open(filepath)
    parts = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            parts.append(text.strip())
    doc.close()

    return {
        "filename": os.path.basename(filepath),
        "text": "\n".join(parts),
        "source_type": "pdf",
    }


def load_document(filepath: str) -> Optional[Dict]:
    """
    根据文件后缀名自动选择解析器。
    返回 {"filename": ..., "text": ..., "source_type": ...}
    解析失败返回 {"error": str}
    """
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower()

    if ext not in (".docx", ".pdf", ".txt"):
        return {"error": f"不支持的文件格式 '{ext}'，仅支持 .docx / .pdf / .txt",
                "filename": filename}

    try:
        if ext == ".docx":
            result = load_word(filepath)
        elif ext == ".pdf":
            result = load_pdf(filepath)
        elif ext == ".txt":
            result = load_txt(filepath)
    except Exception as e:
        error_msg = str(e)[:200]
        # 给出针对性提示
        if ext == ".docx" and "not a valid" in error_msg.lower():
            error_msg = "文件可能不是有效的 .docx 格式（旧版 .doc 文件不支持，请另存为 .docx 后再导入）"
        elif ext == ".pdf" and "no text" in error_msg.lower():
            error_msg = "PDF 文件可能是扫描图片，无法提取文字，请使用 OCR 工具转换后导入"
        return {"error": error_msg, "filename": filename}

    # 校验解析质量
    is_garbled, ratio = _is_text_garbled(result.get("text", ""))
    if is_garbled:
        return {"error": f"解析出的文本可能包含乱码（有效字符比例 {ratio:.0%}），请将文件转换为 TXT 格式后重新导入",
                "filename": filename}

    return result


def load_documents_from_dir(directory: str) -> list:
    """
    从目录批量加载所有支持的文档
    返回成功加载的文档列表，跳过失败文件
    """
    supported_exts = {".docx", ".pdf", ".txt"}
    results = []
    errors = []

    if not os.path.isdir(directory):
        return results, errors

    for filename in sorted(os.listdir(directory)):
        filepath = os.path.join(directory, filename)
        if not os.path.isfile(filepath):
            continue
        ext = os.path.splitext(filename)[1].lower()
        if ext not in supported_exts:
            continue

        doc = load_document(filepath)
        if doc is not None and "error" not in doc:
            results.append(doc)
        else:
            errors.append(filename)

    return results, errors
